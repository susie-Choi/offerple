"""Convert enhanced markdown report to DOCX format."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def process_markdown_line(doc, line, in_code_block, code_lines, in_table, table_lines):
    """Process a single markdown line and add to document."""
    
    # Code block toggle
    if line.startswith('```'):
        if not in_code_block:
            return True, code_lines, in_table, table_lines
        else:
            # End code block
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'Normal'
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            return False, [], in_table, table_lines
    
    if in_code_block:
        code_lines.append(line)
        return in_code_block, code_lines, in_table, table_lines
    
    # Table handling
    if '|' in line and line.strip().startswith('|'):
        table_lines.append(line)
        return in_code_block, code_lines, True, table_lines
    elif in_table and table_lines:
        # Process accumulated table
        process_table(doc, table_lines)
        in_table = False
        table_lines = []
    
    # Headers
    if line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    # Horizontal rule
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.add_run('_' * 80)
    # Bold text
    elif '**' in line:
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
    # Bullet points
    elif line.strip().startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    # Numbered lists
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        doc.add_paragraph(text, style='List Number')
    # Empty line
    elif line.strip() == '':
        doc.add_paragraph()
    # Regular paragraph
    else:
        if line.strip():
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
    
    return in_code_block, code_lines, in_table, table_lines


def process_inline_formatting(paragraph, text):
    """Process inline formatting like bold, code."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def process_table(doc, table_lines):
    """Process markdown table."""
    table_lines = [line for line in table_lines if line.strip() and not line.strip().startswith('|---')]
    
    if len(table_lines) < 2:
        return
    
    # Parse header
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    data_lines = table_lines[1:]
    
    # Create table
    table = doc.add_table(rows=1 + len(data_lines), cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    # Add header
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for row_idx, line in enumerate(data_lines):
        cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
        for col_idx, cell_text in enumerate(cells_data):
            if col_idx < len(table.rows[row_idx + 1].cells):
                table.rows[row_idx + 1].cells[col_idx].text = cell_text


def convert_md_to_docx(md_file, output_file):
    """Main conversion function."""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.rstrip('\n')
        in_code_block, code_lines, in_table, table_lines = process_markdown_line(
            doc, line, in_code_block, code_lines, in_table, table_lines
        )
    
    # Process any remaining table
    if in_table and table_lines:
        process_table(doc, table_lines)
    
    doc.save(output_file)
    print(f"✅ Successfully converted to: {output_file}")


if __name__ == '__main__':
    md_file = Path('docs/nlp-midterm-report-enhanced.md')
    output_file = Path('docs/nlp-midterm-report-enhanced.docx')
    
    if not md_file.exists():
        print(f"❌ Error: {md_file} not found")
        exit(1)
    
    convert_md_to_docx(md_file, output_file)
