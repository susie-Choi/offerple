"""
연구계획서를 Markdown에서 DOCX로 변환하는 스크립트
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_research_proposal_docx():
    """연구계획서 DOCX 파일 생성"""
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    
    # 제목
    title = doc.add_heading('Zero-Day Defense: LLM 기반 잠재 위협 사전 탐지 시스템 연구계획서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 1. 연구 주제
    doc.add_heading('1. 연구 주제', 1)
    
    doc.add_heading('1.1 연구 제목', 2)
    p = doc.add_paragraph()
    p.add_run('LLM 기반 사전 신호 분석을 통한 소프트웨어 생태계 잠재 취약점 예측 시스템').bold = True
    
    doc.add_heading('1.2 연구 배경 및 필요성', 2)
    doc.add_paragraph(
        '현재 소프트웨어 보안은 CVE(Common Vulnerabilities and Exposures) 공개 후 패치를 적용하는 '
        '사후 대응(Reactive) 방식에 의존하고 있습니다. 그러나 Log4Shell(CVE-2021-44228), '
        'Equifax 데이터 유출(CVE-2017-5638) 같은 대규모 보안 사고는 CVE 공개부터 패치 적용까지의 '
        '시간적 공백 동안 막대한 피해를 발생시킵니다.'
    )
    doc.add_paragraph(
        '본 연구는 이러한 한계를 극복하기 위해 CVE 정보가 없는 상태에서 관찰 가능한 '
        '사전 신호(Early Signals)를 분석하여 잠재적 취약점을 예측하는 새로운 패러다임을 제시합니다. '
        '이는 제로데이 공격에 대한 선제적 방어를 가능하게 하며, 제한된 보안 자원을 가장 치명적인 '
        '위협에 집중 배치할 수 있게 합니다.'
    )
    
    doc.add_heading('1.3 연구 목표', 2)
    objectives = [
        '사전 방어 패러다임 구축: CVE 공개 전 잠재 위험 지대(Latent Risk Zone) 식별',
        '다차원 신호 통합: 그래프 구조, 코드 패턴, 개발 활동, 과거 이력 등 다양한 신호 융합',
        'LLM 유추 추론: 과거 취약점 패턴에서 미래 위협을 유추하는 AI 시스템 개발',
        '실세계 검증: Log4Shell, Equifax 등 실제 사고를 사전에 탐지할 수 있었는지 Historical Validation',
        '학술적 기여: NeurIPS, ICML, ICLR 등 top-tier AI 학회 논문 투고'
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f'{i}. {obj}', style='List Number')
    
    doc.add_heading('1.4 연구의 독창성', 2)
    innovations = [
        '패러다임 전환: 사후 대응(Reactive) → 사전 방어(Proactive)',
        '다차원 융합: 그래프 분석 + 코드 분석 + LLM 추론의 하이브리드 시스템',
        '시간적 엄밀성: Data Leakage를 철저히 방지한 Historical Validation',
        '실용적 가치: 전체 패키지의 1-5%만 모니터링하여 대부분의 Critical 취약점 탐지'
    ]
    for innovation in innovations:
        doc.add_paragraph(innovation, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. 연구에 활용할 데이터
    doc.add_heading('2. 연구에 활용할 데이터', 1)
    
    doc.add_heading('2.1 데이터 수집 원칙', 2)
    p = doc.add_paragraph()
    p.add_run('시간적 제약: ').bold = True
    p.add_run('특정 시점 t 이전의 데이터만 사용하여 미래 정보 누락(Data Leakage) 방지')
    
    doc.add_heading('2.2 패키지 생태계 데이터', 2)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('PyPI, Maven Central, npm registry, Libraries.io')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목:').bold = True
    items = [
        '패키지 메타데이터 (이름, 버전, 설명, 라이선스, 작성자)',
        '의존성 관계 (A depends on B)',
        '다운로드 통계 (일별/월별 다운로드 수)',
        '패키지 크기, 파일 수, 릴리스 빈도'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('규모: ').bold = True
    p.add_run('약 10만+ 패키지 (PyPI 기준)')
    
    doc.add_heading('2.3 과거 취약점 이력 데이터', 2)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('NVD (National Vulnerability Database), GitHub Advisory Database, OSS Index, Snyk Vulnerability DB')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목:').bold = True
    items = [
        'CVE ID 및 공개 날짜',
        'CVSS 점수 및 심각도 (Critical, High, Medium, Low)',
        '취약점 유형 (RCE, XSS, SQL Injection, Memory Corruption 등)',
        '영향받는 버전 범위',
        '패치 정보 및 해결 시간'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('중요 제약: ').bold = True
    p.add_run('특정 시점 t 이전에 공개된 CVE만 사용 (예: Log4Shell 검증 시 2021년 11월 1일 이전 데이터만 사용)')
    
    doc.add_heading('2.4 코드 레벨 데이터', 2)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('GitHub 공개 저장소')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목:').bold = True
    items = [
        '소스코드 (Python, Java, JavaScript 등)',
        '위험한 함수 사용 빈도 (strcpy, eval, exec, deserialize, system 등)',
        '코드 복잡도 (Cyclomatic Complexity, Lines of Code)',
        '코드 변경 이력 (커밋 로그, diff 분석)',
        '보안 관련 코드 영역 (네트워크 입력 처리, 파일 업로드, 인증/인가)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('2.5 Historical Validation 데이터', 2)
    p = doc.add_paragraph()
    p.add_run('검증 대상 사건:').bold = True
    
    doc.add_paragraph(
        '1. Log4Shell (CVE-2021-44228): 2021년 12월 공개\n'
        '   - Cutoff: 2021년 11월 1일\n'
        '   - 목표: CVE 공개 1개월 전에 Log4j를 상위 1% 고위험으로 식별'
    )
    doc.add_paragraph(
        '2. Equifax (CVE-2017-5638): 2017년 3월 공개\n'
        '   - Cutoff: 2017년 2월 1일\n'
        '   - 목표: Struts2를 지속적으로 상위 5% 고위험으로 분류'
    )
    doc.add_paragraph(
        '3. 기타 Critical CVE: 2015-2024년 CVSS 9.0+ 취약점 20-30건'
    )
    
    doc.add_page_break()
    
    # 3. 분석하고자 하는 방향
    doc.add_heading('3. 분석하고자 하는 방향', 1)
    
    doc.add_heading('3.1 연구 방법론: 4단계 프레임워크', 2)
    
    doc.add_heading('Phase 1: 핵심 위협 식별 및 중요도 분석', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('오픈소스 생태계에서 "붕괴 시 파급력이 가장 큰" 핵심 노드(Critical Nodes) 식별')
    
    p = doc.add_paragraph()
    p.add_run('분석 방법:').bold = True
    methods = [
        '그래프 중심성 분석: PageRank, Betweenness Centrality, Closeness Centrality 계산',
        '하류 영향 범위: 각 패키지에 의존하는 하위 패키지 수 측정 (Downstream Impact)',
        '과거 취약점 이력: CVSS Critical 등급 취약점 발생 빈도 분석',
        '생태계 중요도: 다운로드 수, 사용 범위, 대체 가능성 평가'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('Phase 2: 위협 군집 분석 및 연관성 탐색', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('핵심 노드들의 숨겨진 연관성을 찾아 "위협 군집(Threat Cluster)" 도출')
    
    p = doc.add_paragraph()
    p.add_run('분석 방법:').bold = True
    methods = [
        '커뮤니티 탐지: Louvain, Label Propagation 알고리즘으로 그래프 군집화',
        '코드 유사도: Code Embedding (CodeBERT, GraphCodeBERT)으로 코드 패턴 유사도 계산',
        '공통 의존성: 동일한 라이브러리를 사용하는 패키지 그룹화',
        '취약점 패턴 유사도: 과거 취약점 유형의 유사성 분석'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('Phase 3: LLM 기반 사전 신호 분석 (핵심 독창성)', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('CVE 공개 전에 관찰 가능한 "사전 신호(Early Signals)"를 통해 잠재 위험 예측')
    
    p = doc.add_paragraph()
    p.add_run('사용 가능한 신호 (CVE 정보 없이 관찰 가능):').bold = True
    
    doc.add_paragraph('1. 과거 취약점 패턴 신호')
    signals = [
        '해당 패키지의 과거 취약점 유형 및 빈도',
        '유사 군집 내 다른 패키지들의 취약점 이력',
        '취약점 발생 간격 및 추세'
    ]
    for signal in signals:
        doc.add_paragraph(signal, style='List Bullet 2')
    
    doc.add_paragraph('2. 코드 레벨 위험 신호')
    signals = [
        '위험한 함수 사용 빈도 (strcpy, eval, exec, deserialize 등)',
        '코드 복잡도 (Cyclomatic Complexity)',
        '최근 대규모 코드 변경 (특히 보안 관련 영역)'
    ]
    for signal in signals:
        doc.add_paragraph(signal, style='List Bullet 2')
    
    doc.add_paragraph('3. 유지보수 활동 신호')
    signals = [
        '보안 패치 빈도 및 속도',
        '개발자 활동 패턴 (활발함 vs 방치)',
        '이슈 트래커의 보안 관련 논의'
    ]
    for signal in signals:
        doc.add_paragraph(signal, style='List Bullet 2')
    
    doc.add_paragraph('4. 커뮤니티 및 생태계 신호')
    signals = [
        '급격한 다운로드 수 증가 (공격자의 관심 증가 가능)',
        '포크 수 급증 (보안 문제 우회 시도 가능)',
        '관련 패키지들의 취약점 발생 추세'
    ]
    for signal in signals:
        doc.add_paragraph(signal, style='List Bullet 2')
    
    p = doc.add_paragraph()
    p.add_run('LLM 유추 추론: ').bold = True
    p.add_run(
        '각 패키지의 다차원 신호를 컨텍스트로 제공하여 LLM이 과거 패턴에서 미래 위협을 유추. '
        '예: "패키지 A는 과거 메모리 누수 취약점 3회 + libX 사용 + 최근 메모리 관리 코드 변경 → '
        '동일 군집의 B, C도 libX 사용 + 유사 코드 패턴 → 잠재 위험 높음"'
    )
    
    doc.add_heading('Phase 4: 3D 시각화 및 방어 전략 수립', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('잠재 위험 지대를 직관적으로 시각화하여 자원 배치 최적화')
    
    p = doc.add_paragraph()
    p.add_run('분석 방법:').bold = True
    methods = [
        '차원 축소: t-SNE 또는 UMAP으로 고차원 신호를 2D로 축소',
        '3D 매핑: X, Y축은 패키지 간 유사도, Z축은 잠재 위험 점수',
        '위험 지대 식별: Z축이 높은 영역 = 잠재 위험 지대',
        '인터랙티브 시각화: Plotly 기반 3D 그래프'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('3.2 평가 방법론', 2)
    
    doc.add_heading('Historical Validation (핵심 검증 방법)', 3)
    p = doc.add_paragraph()
    p.add_run('원칙: ').bold = True
    p.add_run('특정 시점 t 이전의 데이터만 사용하여 t 이후 발생한 취약점 예측')
    
    p = doc.add_paragraph()
    p.add_run('평가 지표:').bold = True
    metrics = [
        'Precision@K: 상위 K개 예측 중 실제로 취약점이 발생한 비율',
        'Recall@K: 실제 발생한 Critical/High 취약점 중 상위 K개에 포함된 비율',
        'Lead Time: CVE 공개 시점 대비 얼마나 일찍 위험 신호를 탐지했는지',
        'False Positive Rate: 잘못된 경보 비율'
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('목표 성능:').bold = True
    goals = [
        'Precision@100 ≥ 20% (상위 1% 중 20% 이상이 실제 취약점 발생)',
        'Recall@100 ≥ 50% (Critical CVE의 50% 이상을 상위 1%에서 탐지)',
        'Lead Time ≥ 30일 (CVE 공개 1개월 전에 탐지)'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_heading('3.3 기대 성과', 2)
    
    p = doc.add_paragraph()
    p.add_run('학술적 기여:').bold = True
    contributions = [
        '새로운 연구 분야: Zero-Day Defense라는 새로운 연구 방향 개척',
        '방법론적 혁신: 그래프 분석 + LLM 유추 추론의 하이브리드 시스템',
        '벤치마크 구축: 잠재 위협 탐지 성능을 평가할 수 있는 새로운 벤치마크',
        '실세계 검증: Log4Shell, Equifax 등 실제 사고로 검증'
    ]
    for contrib in contributions:
        doc.add_paragraph(contrib, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('실용적 기여:').bold = True
    contributions = [
        '자원 최적화: 전체 패키지의 1-5%만 모니터링하여 대부분의 Critical 취약점 탐지',
        '조기 경보: CVE 공개 전 잠재 위험 지대 식별',
        '의사결정 지원: 보안 담당자가 어디에 자원을 집중할지 결정',
        '오픈소스 도구: 연구 결과를 오픈소스로 공개하여 커뮤니티 기여'
    ]
    for contrib in contributions:
        doc.add_paragraph(contrib, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('사회적 임팩트:').bold = True
    impacts = [
        '경제적 가치: Log4Shell(10년 복구 비용), Equifax(14억 달러) 같은 피해 사전 예방',
        '패러다임 전환: "제로데이는 막을 수 없다"는 고정관념을 깨는 새로운 접근',
        '보안 민주화: 대기업뿐 아니라 중소기업도 사용할 수 있는 도구 제공'
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style='List Bullet')
    
    # 문서 저장
    output_path = '연구계획서.docx'
    doc.save(output_path)
    print(f"연구계획서가 생성되었습니다: {output_path}")
    return output_path

if __name__ == "__main__":
    create_research_proposal_docx()
