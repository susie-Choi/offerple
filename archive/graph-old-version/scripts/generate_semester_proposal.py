"""
한 학기 연구계획서를 DOCX로 변환하는 스크립트
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_border(table):
    """테이블에 테두리 추가"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_semester_proposal_docx():
    """한 학기 연구계획서 DOCX 파일 생성"""
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    
    # 제목
    title = doc.add_heading('Zero-Day Defense 연구 프로젝트 계획서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('(한 학기 버전)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 1. 연구 주제
    doc.add_heading('1. 연구 주제', 1)
    
    doc.add_heading('1.1 연구 제목', 2)
    p = doc.add_paragraph()
    p.add_run('그래프 분석과 LLM을 활용한 오픈소스 패키지 취약점 위험도 예측 시스템 프로토타입').bold = True
    
    doc.add_heading('1.2 연구 배경', 2)
    doc.add_paragraph(
        '현재 소프트웨어 보안은 CVE(Common Vulnerabilities and Exposures) 공개 후 패치를 적용하는 '
        '사후 대응 방식에 의존하고 있습니다. Log4Shell(CVE-2021-44228) 같은 대규모 보안 사고는 '
        'CVE 공개부터 패치 적용까지의 시간적 공백 동안 막대한 피해를 발생시킵니다.'
    )
    doc.add_paragraph(
        '본 연구는 CVE 정보가 없는 상태에서 관찰 가능한 사전 신호(그래프 구조, 과거 이력, 코드 패턴)를 '
        '분석하여 잠재적으로 위험한 패키지를 예측하는 시스템의 프로토타입을 개발합니다.'
    )
    
    doc.add_heading('1.3 한 학기 연구 목표', 2)
    objectives = [
        '소규모 데이터셋 구축: Python 생태계(PyPI)의 주요 패키지 1,000-5,000개 대상',
        '기본 신호 추출: 그래프 중심성, 과거 취약점 이력, 기본 코드 메트릭',
        '간단한 위험도 예측 모델: 신호 기반 점수 계산 + LLM 보조 분석',
        '1개 사례 검증: Log4Shell을 Historical Validation으로 검증',
        '프로토타입 시스템: 기본적인 API 및 시각화 인터페이스'
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f'{i}. {obj}', style='List Number')
    
    doc.add_heading('1.4 연구 범위 제한 (한 학기 내 달성 가능)', 2)
    limitations = [
        '데이터 규모: 전체 생태계가 아닌 주요 패키지 1,000-5,000개로 제한',
        '신호 종류: 핵심 신호 3-4가지로 제한 (그래프, 과거 이력, 기본 코드 메트릭)',
        '검증 사례: Log4Shell 1건으로 제한 (Equifax 등은 향후 연구로)',
        'LLM 활용: 전체 패키지가 아닌 상위 위험 후보 100개만 분석',
        '시각화: 3D 대신 2D 시각화로 단순화'
    ]
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. 연구에 활용할 데이터
    doc.add_heading('2. 연구에 활용할 데이터', 1)
    
    doc.add_heading('2.1 데이터 수집 범위 (현실적 규모)', 2)
    
    doc.add_heading('2.1.1 패키지 선정 기준', 3)
    p = doc.add_paragraph()
    p.add_run('대상: ').bold = True
    p.add_run('PyPI 생태계의 주요 패키지 1,000-5,000개')
    
    p = doc.add_paragraph()
    p.add_run('선정 기준:').bold = True
    criteria = [
        '다운로드 수 상위 1,000개 (인기 패키지)',
        '과거 CVE 이력이 있는 패키지 500개',
        '핵심 인프라 패키지 (requests, numpy, django 등) 500개',
        '무작위 샘플 1,000-3,000개'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('예상 데이터 규모: ').bold = True
    p.add_run('약 5-10GB (전체 생태계 대비 1/10 규모)')
    
    doc.add_heading('2.1.2 패키지 메타데이터', 3)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('PyPI API (https://pypi.org/pypi/{package}/json)')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목:').bold = True
    items = [
        '패키지 이름, 버전, 설명',
        '의존성 관계 (requires_dist)',
        '다운로드 통계 (pypistats API)',
        '릴리스 날짜, 유지보수 상태'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('2.1.3 과거 취약점 이력', 3)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('NVD (National Vulnerability Database) API, GitHub Advisory Database')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목:').bold = True
    items = [
        'CVE ID 및 공개 날짜',
        'CVSS 점수 (Critical/High만 집중)',
        '취약점 유형 (RCE, XSS 등)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('시간 범위: ').bold = True
    p.add_run('2015-2024년 (10년간)')
    
    p = doc.add_paragraph()
    p.add_run('중요 제약: ').bold = True
    p.add_run('Log4Shell 검증 시 2021년 11월 1일 이전 데이터만 사용')
    
    doc.add_heading('2.1.4 코드 레벨 데이터 (간소화)', 3)
    p = doc.add_paragraph()
    p.add_run('출처: ').bold = True
    p.add_run('GitHub 공개 저장소')
    
    p = doc.add_paragraph()
    p.add_run('수집 항목 (기본 메트릭만):').bold = True
    items = [
        '코드 복잡도 (radon 라이브러리 사용)',
        '위험 함수 키워드 검색 (eval, exec, system 등)',
        '파일 수, 코드 라인 수'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('제한사항: ').bold = True
    p.add_run('전체 소스코드 다운로드 대신 GitHub API로 통계만 수집, 상위 위험 후보 100개만 상세 분석')
    
    doc.add_page_break()
    
    # 3. 분석 방향 및 단계별 계획
    doc.add_heading('3. 분석 방향 및 단계별 계획', 1)
    
    doc.add_heading('3.1 전체 연구 흐름 (3단계로 단순화)', 2)
    doc.add_paragraph('[Phase 1] 데이터 수집 및 그래프 구축 (3주)')
    doc.add_paragraph('    ↓')
    doc.add_paragraph('[Phase 2] 신호 추출 및 위험도 점수 계산 (4주)')
    doc.add_paragraph('    ↓')
    doc.add_paragraph('[Phase 3] Log4Shell 검증 및 시각화 (3주)')
    
    doc.add_heading('3.2 Phase 1: 데이터 수집 및 그래프 구축 (3주)', 2)
    
    doc.add_heading('Week 1-2: 패키지 데이터 수집', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('PyPI에서 1,000-5,000개 패키지 메타데이터 수집')
    
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        'PyPI API로 인기 패키지 목록 수집',
        '각 패키지의 메타데이터 다운로드',
        '의존성 관계 파싱',
        'JSON 파일로 저장'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('산출물: ').bold = True
    p.add_run('packages.json (패키지 메타데이터), dependencies.json (의존성 관계)')
    
    doc.add_heading('Week 3: 취약점 이력 수집 및 그래프 구축', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('NVD에서 CVE 데이터 수집, 의존성 그래프 생성')
    
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        'NVD API로 2015-2024년 CVE 데이터 수집',
        'CVE와 패키지 매칭',
        'NetworkX로 의존성 그래프 생성',
        '그래프 기본 통계 확인 (노드 수, 엣지 수, 연결 성분)'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('산출물: ').bold = True
    p.add_run('vulnerabilities.json, dependency_graph.gpickle, graph_stats.txt')
    
    doc.add_heading('3.3 Phase 2: 신호 추출 및 위험도 점수 계산 (4주)', 2)
    
    doc.add_heading('Week 4: 그래프 신호 추출', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('그래프 중심성 지표 계산')
    
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        'PageRank 계산 (전역 중요도)',
        'In-degree 계산 (얼마나 많은 패키지가 의존하는지)',
        'Downstream Impact 계산 (하류 영향 범위)'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    doc.add_heading('Week 5: 과거 취약점 패턴 신호 추출', 3)
    p = doc.add_paragraph()
    p.add_run('작업 내용: ').bold = True
    p.add_run('패키지별 CVE 빈도, Critical/High CVE 비율, 평균 CVSS 점수, 최근 취약점 발생 추세 분석')
    
    doc.add_heading('Week 6: 코드 레벨 신호 추출 (간소화)', 3)
    p = doc.add_paragraph()
    p.add_run('작업 내용: ').bold = True
    p.add_run('상위 위험 후보 100개의 GitHub 통계 수집, 위험 키워드 검색, 코드 복잡도 계산')
    
    doc.add_heading('Week 7: 위험도 점수 계산 및 LLM 분석', 3)
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        '신호 정규화 (0-1 범위)',
        '가중 평균으로 기본 위험도 점수 계산',
        '상위 100개 패키지에 대해 LLM 분석 (GPT-4 또는 Claude)',
        '최종 위험도 점수 계산'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('산출물: ').bold = True
    p.add_run('risk_scores.json, llm_analysis.json')
    
    doc.add_page_break()
    
    doc.add_heading('3.4 Phase 3: Log4Shell 검증 및 시각화 (3주)', 2)
    
    doc.add_heading('Week 8-9: Log4Shell Historical Validation', 3)
    p = doc.add_paragraph()
    p.add_run('목표: ').bold = True
    p.add_run('2021년 11월 1일 시점에 Log4j를 고위험으로 식별할 수 있었는지 검증')
    
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        '시간 분할 설정 (Cutoff: 2021-11-01, CVE 공개: 2021-12-09)',
        '과거 데이터 재구성 (2021년 11월 1일 이전 데이터만 사용)',
        '예측 실행 (Phase 2의 전체 파이프라인)',
        '결과 분석 (Log4j의 위험도 순위, 주요 신호, Lead Time 계산)'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('성공 기준: ').bold = True
    p.add_run('Log4j가 상위 5% 이내에 포함되면 성공, 상위 1% 이내면 매우 성공적')
    
    doc.add_heading('Week 10: 시각화 및 프로토타입 완성', 3)
    p = doc.add_paragraph()
    p.add_run('작업 내용:').bold = True
    tasks = [
        '2D 위험 지도 생성 (X축: PageRank, Y축: 과거 취약점 빈도)',
        '간단한 웹 인터페이스 구축 (Streamlit 사용)',
        '보고서 작성 (연구 방법론, Log4Shell 검증 결과, 한계점)'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('산출물: ').bold = True
    p.add_run('risk_map_2d.html, Streamlit 앱, final_report.pdf')
    
    doc.add_page_break()
    
    # 4. 한 학기 일정표
    doc.add_heading('4. 한 학기 일정표', 1)
    
    # 테이블 생성
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # 헤더
    header_cells = table.rows[0].cells
    header_cells[0].text = '주차'
    header_cells[1].text = '단계'
    header_cells[2].text = '주요 작업'
    header_cells[3].text = '산출물'
    
    # 데이터
    schedule_data = [
        ('1-2', 'Phase 1', '패키지 메타데이터 수집', 'packages.json, dependencies.json'),
        ('3', 'Phase 1', 'CVE 수집 및 그래프 구축', 'vulnerabilities.json, dependency_graph.gpickle'),
        ('4', 'Phase 2', '그래프 신호 추출', 'graph_signals.json'),
        ('5', 'Phase 2', '취약점 패턴 신호 추출', 'vulnerability_patterns.json'),
        ('6', 'Phase 2', '코드 신호 추출 (100개)', 'code_signals.json'),
        ('7', 'Phase 2', '위험도 점수 계산 + LLM', 'risk_scores.json, llm_analysis.json'),
        ('8-9', 'Phase 3', 'Log4Shell 검증', 'log4shell_validation.json'),
        ('10', 'Phase 3', '시각화 및 보고서', 'risk_map_2d.html, final_report.pdf'),
    ]
    
    for i, (week, phase, task, output) in enumerate(schedule_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = week
        row_cells[1].text = phase
        row_cells[2].text = task
        row_cells[3].text = output
    
    add_table_border(table)
    
    doc.add_paragraph()
    
    # 5. 기대 성과
    doc.add_heading('5. 기대 성과 (한 학기 범위)', 1)
    
    doc.add_heading('5.1 학술적 성과', 2)
    achievements = [
        '프로토타입 시스템: 잠재 위협 탐지의 가능성 입증',
        '방법론 검증: 사전 신호 기반 예측의 유효성 확인',
        '향후 연구 기반: 전체 규모 연구를 위한 기초 마련'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('5.2 실용적 성과', 2)
    achievements = [
        '위험 패키지 식별: 1,000-5,000개 중 고위험 패키지 목록',
        '시각화 도구: 패키지 위험도를 직관적으로 확인할 수 있는 인터페이스',
        '오픈소스 공개: 코드 및 데이터를 GitHub에 공개'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('5.3 학습 성과', 2)
    achievements = [
        '그래프 분석 기법 습득',
        'LLM API 활용 경험',
        '보안 데이터 분석 역량',
        '실세계 문제 해결 경험'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. 필요 리소스
    doc.add_heading('6. 필요 리소스', 1)
    
    doc.add_heading('6.1 개발 환경', 2)
    resources = [
        'Python 3.9+',
        '주요 라이브러리: NetworkX, pandas, requests, openai/anthropic',
        '개발 도구: VS Code, Jupyter Notebook'
    ]
    for resource in resources:
        doc.add_paragraph(resource, style='List Bullet')
    
    doc.add_heading('6.2 API 및 서비스', 2)
    services = [
        'PyPI API (무료)',
        'NVD API (무료, rate limit 있음)',
        'GitHub API (무료, rate limit 있음)',
        'OpenAI API 또는 Claude API (유료, 예상 비용: $50-100)'
    ]
    for service in services:
        doc.add_paragraph(service, style='List Bullet')
    
    doc.add_heading('6.3 컴퓨팅 리소스', 2)
    resources = [
        '로컬 PC (8GB RAM 이상 권장)',
        '저장 공간: 20GB 이상'
    ]
    for resource in resources:
        doc.add_paragraph(resource, style='List Bullet')
    
    # 7. 위험 요소 및 대응 방안
    doc.add_heading('7. 위험 요소 및 대응 방안', 1)
    
    doc.add_heading('7.1 데이터 수집 지연', 2)
    p = doc.add_paragraph()
    p.add_run('위험: ').bold = True
    p.add_run('API rate limit으로 데이터 수집 지연')
    p = doc.add_paragraph()
    p.add_run('대응: ').bold = True
    p.add_run('캐싱 활용, 필요시 패키지 수를 1,000개로 축소, 미리 수집된 데이터셋 활용')
    
    doc.add_heading('7.2 LLM API 비용', 2)
    p = doc.add_paragraph()
    p.add_run('위험: ').bold = True
    p.add_run('LLM API 비용 초과')
    p = doc.add_paragraph()
    p.add_run('대응: ').bold = True
    p.add_run('분석 대상을 상위 50개로 축소, GPT-3.5 같은 저렴한 모델 사용, 배치 처리로 비용 절감')
    
    doc.add_heading('7.3 Log4Shell 검증 실패', 2)
    p = doc.add_paragraph()
    p.add_run('위험: ').bold = True
    p.add_run('Log4j를 고위험으로 식별하지 못함')
    p = doc.add_paragraph()
    p.add_run('대응: ').bold = True
    p.add_run('신호 가중치 조정, 추가 신호 탐색, 실패 원인 분석도 중요한 연구 결과')
    
    # 8. 향후 확장 방향
    doc.add_heading('8. 향후 확장 방향 (한 학기 이후)', 1)
    
    doc.add_heading('8.1 단기 확장 (방학 중)', 2)
    expansions = [
        '패키지 수 확대: 10,000-50,000개',
        '추가 검증 사례: Equifax, 기타 Critical CVE',
        '코드 분석 강화: 전체 패키지 대상'
    ]
    for expansion in expansions:
        doc.add_paragraph(expansion, style='List Bullet')
    
    doc.add_heading('8.2 중기 확장 (다음 학기)', 2)
    expansions = [
        '다른 생태계 추가: Maven (Java), npm (JavaScript)',
        '3D 시각화 구현',
        '실시간 모니터링 시스템'
    ]
    for expansion in expansions:
        doc.add_paragraph(expansion, style='List Bullet')
    
    doc.add_heading('8.3 장기 목표 (1년 후)', 2)
    goals = [
        '전체 규모 시스템 구축',
        '학술 논문 작성 및 투고',
        '실용 서비스 런칭'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    # 문서 저장
    output_path = '연구계획서_한학기.docx'
    doc.save(output_path)
    print(f"한 학기 연구계획서가 생성되었습니다: {output_path}")
    return output_path

if __name__ == "__main__":
    create_semester_proposal_docx()
