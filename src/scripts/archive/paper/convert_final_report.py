"""Convert final markdown report to DOCX format."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def process_inline_formatting(paragraph, text):
    """Process inline formatting like bold, italic, code."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def process_table(doc, table_lines):
    """Process markdown table."""
    table_lines = [line for line in table_lines if line.strip() and not line.strip().startswith('|---')]
    
    if len(table_lines) < 2:
        return
    
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    data_lines = table_lines[1:]
    
    table = doc.add_table(rows=1 + len(data_lines), cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for row_idx, line in enumerate(data_lines):
        cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
        for col_idx, cell_text in enumerate(cells_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text

def parse_markdown_to_docx(md_file, output_file):
    """Convert markdown to DOCX."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Normal'
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            if i >= len(lines) or '|' not in lines[i]:
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
        elif '**' in line:
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        elif line.strip() == '':
            doc.add_paragraph()
        else:
            if line.strip():
                p = doc.add_paragraph()
                process_inline_formatting(p, line)
        
        i += 1
    
    doc.save(output_file)
    print(f"✅ Successfully converted to: {output_file}")

if __name__ == '__main__':
    md_file = Path('docs/nlp-midterm-report-final.md')
    output_file = Path('docs/nlp-midterm-report-final.docx')
    
    if not md_file.exists():
        print(f"❌ Error: {md_file} not found")
        exit(1)
    
    parse_markdown_to_docx(md_file, output_file)
